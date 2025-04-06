from docx import Document
from docx.shared import Inches

# Create a new Word document
doc = Document()
doc.add_heading('SuperPower Labs Server Infrastructure Plan', 0)

# Summary section
doc.add_heading('Project Summary', level=1)
doc.add_paragraph(
    "This document outlines the infrastructure and planning for SuperPower Labs’ self-hosted server network. "
    "The system spans Hostinger cloud services and a home network, featuring Dockerized applications, native installs, "
    "automated backups, monitoring, and secure connectivity via Tailscale."
)

# VPS Specs
doc.add_heading('Server Specifications', level=1)
doc.add_paragraph("- KVM8 (x2): 8 vCPU, 32GB RAM, 400GB NVMe, 32TB bandwidth")
doc.add_paragraph("- KVM2 (x1): 2 vCPU, 8GB RAM, 100GB NVMe, 8TB bandwidth")
doc.add_paragraph("- The Beast (Workstation): Specs TBD")

# Network Overview
doc.add_heading('Infrastructure Overview by Network', level=1)

doc.add_heading('Hostinger Network', level=2)
doc.add_paragraph("- Cloudflare DNS and Hostinger firewall settings per VPS")
doc.add_paragraph("- KVM8-A: BigBlueButton (native), Docker, Portainer, Tailscale, Cockpit")
doc.add_paragraph("- KVM8-B: PostgreSQL, Moodle, n8n, NGINX Proxy Manager, Portainer, Docker, Tailscale, Cockpit")
doc.add_paragraph("- KVM2: Monitoring (Grafana, Prometheus), Dashy/Homer, Docker, Portainer, Cockpit, Tailscale")

doc.add_heading('Home Network', level=2)
doc.add_paragraph("- AT&T Fiber Gateway (1Gbps up/down)")
doc.add_paragraph("- Synology NAS: Backup destination (rsync or Active Backup)")
doc.add_paragraph("- Optional Raspberry Pis for media server, DNS, automation")
doc.add_paragraph("- The Beast: Local workstation, LAN connected")

# Key Phases
doc.add_heading('Setup Phases', level=1)
phases = [
    "Phase 1: OS & Security - SSH keys, firewall, Tailscale, Cockpit",
    "Phase 2: Docker Infrastructure - Docker, Portainer, folder structure",
    "Phase 3: Reverse Proxy - NGINX Proxy Manager with Cloudflare",
    "Phase 4: App Deployment - BBB (native), Moodle, PostgreSQL, n8n (Docker)",
    "Phase 5: AI Stack - LLMs (Ollama, LLaMA) on KVM8-B or KVM2",
    "Phase 6: Monitoring/Backups - Netdata, Watchtower, rsync to Synology",
    "Phase 7: Scalability - Portable Docker, `.env` secrets, optional Docker Swarm"
]
for item in phases:
    doc.add_paragraph(f"- {item}", style='List Bullet')

# Risk Planning
doc.add_heading('Risk Areas & Planning', level=1)
risks = [
    "Permission issues (UID/GID mismatches)",
    "Reverse proxy config and SSL cert delays",
    "BigBlueButton port conflicts",
    "Container dependency timing (Postgres before Moodle)",
    "DNS propagation and SSL failures",
    "Resource contention (BBB and LLMs need isolation)"
]
for risk in risks:
    doc.add_paragraph(f"- {risk}", style='List Bullet')

# Save the document
output_path = "/mnt/data/SuperPowerLabs_ServerPlan.docx"
doc.save(output_path)

output_path
